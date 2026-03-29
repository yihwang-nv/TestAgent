"""
Persistent vector knowledge base for claude_agent.py (RAG + long-term notes).

Uses ChromaDB + sentence-transformers. PDF via PyMuPDF; optional Tesseract OCR for
scanned PDFs. .docx via python-docx.
"""
from __future__ import annotations

import hashlib
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

# PDF OCR: auto | always | never
_PDF_OCR_MODE = os.environ.get("AGENT_KB_PDF_OCR", "auto").strip().lower()
# Tesseract language pack(s), e.g. eng, chi_sim, chi_sim+eng
_OCR_LANG = os.environ.get("AGENT_KB_OCR_LANG", "chi_sim+eng").strip() or "eng"


def _ocr_max_pages() -> Optional[int]:
    """Max PDF pages to OCR; <=0 or unset 0 means no limit (use with care). Default 100."""
    raw = os.environ.get("AGENT_KB_OCR_MAX_PAGES", "100").strip()
    if not raw:
        return 100
    try:
        n = int(raw)
    except ValueError:
        return 100
    if n <= 0:
        return None
    return n

CHUNK_CHARS = 900
CHUNK_OVERLAP = 120

_active_store: Optional["AgentKnowledgeBase"] = None


def kb_deps_available() -> bool:
    try:
        import chromadb  # noqa: F401
        from chromadb.utils import embedding_functions  # noqa: F401
        import sentence_transformers  # noqa: F401
        return True
    except ImportError:
        return False


def kb_install_hint() -> str:
    return (
        "pip install chromadb sentence-transformers pymupdf python-docx "
        "pytesseract pillow  # OCR 还需系统安装 tesseract-ocr（及语言包 chi_sim 等）"
    )


def _load_docx(path: Path) -> tuple[Optional[str], Optional[str]]:
    try:
        import docx
    except ImportError:
        return None, "需要 python-docx: pip install python-docx"
    try:
        d = docx.Document(path)
        parts: list[str] = []
        for p in d.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in d.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [(c.text or "").strip() for c in row.cells]
                if any(cells):
                    rows.append("\t".join(cells))
            if rows:
                parts.append("\n".join(rows))
        text = "\n\n".join(parts)
        if not text.strip():
            return None, "DOCX 中无文本"
        return text, None
    except Exception as e:
        return None, f"读取 DOCX 失败: {e}"


def _pdf_text_per_page(doc: Any) -> list[str]:
    return [page.get_text() or "" for page in doc]


def _ocr_pdf_pages(path: Path, max_pages: Optional[int] = None) -> tuple[Optional[str], Optional[str]]:
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError as e:
        return None, f"OCR 依赖缺失: {e}（pip install pymupdf pytesseract pillow，并安装系统 tesseract）"
    try:
        pytesseract.get_tesseract_version()
    except Exception as e:
        return None, f"未检测到 Tesseract 可执行文件: {e}（如: apt install tesseract-ocr tesseract-ocr-chi-sim）"

    doc = fitz.open(path)
    try:
        n = len(doc)
        limit = n if max_pages is None else min(n, max_pages)
        zoom = float(os.environ.get("AGENT_KB_OCR_ZOOM", "2.0"))
        zoom = max(1.0, min(zoom, 4.0))
        mat = fitz.Matrix(zoom, zoom)
        pieces: list[str] = []
        for i in range(limit):
            page = doc[i]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            try:
                txt = pytesseract.image_to_string(img, lang=_OCR_LANG)
            except Exception as ex:
                txt = ""
                pieces.append(f"[page {i + 1} OCR error: {ex}]")
            t = (txt or "").strip()
            if t:
                pieces.append(f"--- page {i + 1} ---\n{t}")
        merged = "\n\n".join(pieces)
        return (merged if merged.strip() else None), None
    finally:
        doc.close()


def _pdf_native_then_maybe_ocr(path: Path) -> tuple[Optional[str], Optional[str]]:
    try:
        import fitz
    except ImportError:
        return None, "需要 pymupdf: pip install pymupdf"

    doc = fitz.open(path)
    try:
        pages_text = _pdf_text_per_page(doc)
    finally:
        doc.close()

    native = "\n\n".join(pages_text)
    n = len(pages_text)
    stripped = native.strip()
    mode = _PDF_OCR_MODE if _PDF_OCR_MODE in ("auto", "always", "never") else "auto"

    need_ocr = False
    if mode == "always":
        need_ocr = n > 0
    elif mode == "never":
        need_ocr = False
    else:
        if n == 0:
            need_ocr = False
        else:
            avg = len(stripped) / n
            need_ocr = len(stripped) < 80 or avg < 35

    if not need_ocr:
        return native, None

    ocr_text, ocr_err = _ocr_pdf_pages(path, max_pages=_ocr_max_pages())
    if ocr_err and not ocr_text:
        if stripped:
            return native, f"（OCR 未启用或失败，已仅用 PDF 文本层）{ocr_err}"
        return None, ocr_err
    if ocr_text and len(ocr_text.strip()) >= len(stripped):
        return ocr_text, None
    if ocr_text and stripped:
        return f"{native}\n\n--- OCR supplement ---\n{ocr_text}", None
    if ocr_text:
        return ocr_text, None
    if stripped:
        return native, None
    return None, "无文本层且 OCR 无输出（检查扫描质量与 Tesseract 语言包）"


def set_kb_store(store: Optional["AgentKnowledgeBase"]) -> None:
    global _active_store
    _active_store = store


def get_kb_store() -> Optional["AgentKnowledgeBase"]:
    return _active_store


def chunk_text(text: str, size: int = CHUNK_CHARS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    chunks: list[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + size, n)
        chunks.append(text[start:end])
        if end >= n:
            break
        start = max(0, end - overlap)
    return chunks


def _stable_chunk_id(source_key: str, index: int) -> str:
    h = hashlib.sha256(f"{source_key}\0{index}".encode()).hexdigest()
    return h[:40]


def load_file_as_text(path: Path) -> tuple[Optional[str], Optional[str]]:
    if not path.is_file():
        return None, f"Not a file: {path}"
    suf = path.suffix.lower()
    try:
        if suf == ".pdf":
            return _pdf_native_then_maybe_ocr(path)
        if suf == ".docx":
            return _load_docx(path)
        if suf in {
            ".txt",
            ".md",
            ".py",
            ".json",
            ".yaml",
            ".yml",
            ".csv",
            ".rst",
            ".toml",
            ".tsx",
            ".ts",
            ".jsx",
            ".js",
            ".html",
            ".htm",
            ".xml",
            ".c",
            ".h",
            ".cpp",
            ".hpp",
            ".rs",
            ".go",
            ".java",
            ".kt",
            ".swift",
            ".rb",
            ".sh",
            ".sql",
        }:
            return path.read_text(encoding="utf-8", errors="replace"), None
        # try utf-8 for unknown extension (small files only)
        if path.stat().st_size <= 2 * 1024 * 1024:
            return path.read_text(encoding="utf-8", errors="replace"), None
        return None, f"不支持的类型或文件过大: {path.name} ({suf})"
    except Exception as e:
        return None, f"读取失败: {e}"


class AgentKnowledgeBase:
    def __init__(self, persist_dir: Path, embed_model: str = "all-MiniLM-L6-v2") -> None:
        if not kb_deps_available():
            raise ImportError(kb_install_hint())
        self.persist_dir = Path(persist_dir)
        self.persist_dir.mkdir(parents=True, exist_ok=True)
        self.embed_model = embed_model
        self._collection = None

    @property
    def collection(self) -> Any:
        if self._collection is None:
            import chromadb
            from chromadb.utils import embedding_functions

            client = chromadb.PersistentClient(path=str(self.persist_dir))
            ef = embedding_functions.SentenceTransformerEmbeddingFunction(
                model_name=self.embed_model,
            )
            self._collection = client.get_or_create_collection(
                name="claude_agent_kb",
                embedding_function=ef,
                metadata={"hnsw:space": "cosine"},
            )
        return self._collection

    def _delete_by_source(self, source: str) -> None:
        try:
            self.collection.delete(where={"source": source})
        except Exception:
            pass

    def ingest_file(self, path: Path) -> str:
        text, err = load_file_as_text(path)
        if err:
            return f"Error: {err}"
        if not text or not text.strip():
            return f"Error: 无文本内容（可能为扫描版 PDF）: {path}"
        resolved = str(path.resolve())
        title = path.name
        self._delete_by_source(resolved)
        chunks = chunk_text(text)
        ids = [_stable_chunk_id(resolved, i) for i in range(len(chunks))]
        metadatas = [
            {
                "kind": "document",
                "source": resolved,
                "title": title,
                "chunk_index": str(i),
            }
            for i in range(len(chunks))
        ]
        self.collection.upsert(ids=ids, documents=chunks, metadatas=metadatas)
        return (
            f"已入库 {len(chunks)} 条文本块 · 文件: {title} · 路径: {resolved}"
        )

    def remember(self, text: str, title: Optional[str] = None) -> str:
        text = (text or "").strip()
        if not text:
            return "Error: 空内容"
        if len(text) > 200_000:
            return "Error: 内容过长（>200k 字符），请缩短或拆成多次保存"
        now = datetime.now(timezone.utc).isoformat()
        mem_key = f"memory:{hashlib.sha256((title or '') + text[:2000].encode()).hexdigest()[:24]}"
        self._delete_by_source(mem_key)
        chunks = chunk_text(text)
        ids = [_stable_chunk_id(mem_key, i) for i in range(len(chunks))]
        label = (title or "note").strip() or "note"
        metadatas = [
            {
                "kind": "memory",
                "source": mem_key,
                "title": label,
                "created": now,
                "chunk_index": str(i),
            }
            for i in range(len(chunks))
        ]
        self.collection.upsert(ids=ids, documents=chunks, metadatas=metadatas)
        return f"已保存记忆 {len(chunks)} 块 · 标题: {label}"

    def search(self, query: str, top_k: int = 8, kind: Optional[str] = None) -> str:
        query = (query or "").strip()
        if not query:
            return "Error: 空查询"
        try:
            tk = top_k[0] if isinstance(top_k, (list, tuple)) and top_k else top_k
            k = max(1, min(int(tk), 32))
        except (TypeError, ValueError):
            k = 8
        where: Optional[dict[str, str]] = None
        if kind in ("memory", "document"):
            where = {"kind": kind}
        res = self.collection.query(
            query_texts=[query],
            n_results=k,
            where=where,
        )
        docs = (res.get("documents") or [[]])[0]
        metas = (res.get("metadatas") or [[]])[0]
        dists = (res.get("distances") or [[]])[0]
        if not docs:
            return "（知识库中无匹配片段；可先 kb_ingest_file 添加 PDF/文档，或 kb_remember 保存要点）"
        lines: list[str] = []
        for i, doc in enumerate(docs):
            meta = metas[i] if i < len(metas) else {}
            d = dists[i] if i < len(dists) else None
            title = meta.get("title", "?")
            src = meta.get("source", "?")
            kind_m = meta.get("kind", "?")
            dist_s = f"{d:.4f}" if isinstance(d, (int, float)) else "?"
            lines.append(
                f"### [{i + 1}] kind={kind_m} title={title!r} source={src!r} distance={dist_s}\n{doc}"
            )
        return "\n\n".join(lines)

    def context_for_prompt(self, query: str, top_k: int = 5) -> str:
        """Short RAG block to prepend; empty if nothing useful."""
        q = (query or "").strip()
        if not q:
            return ""
        try:
            tk = top_k[0] if isinstance(top_k, (list, tuple)) and top_k else top_k
            nres = max(1, min(int(tk), 10))
        except (TypeError, ValueError):
            nres = 5
        res = self.collection.query(query_texts=[q], n_results=nres)
        docs = (res.get("documents") or [[]])[0]
        metas = (res.get("metadatas") or [[]])[0]
        if not docs:
            return ""
        parts = []
        for doc, meta in zip(docs, metas):
            title = meta.get("title", "")
            parts.append(f"[{meta.get('kind', '?')}: {title}]\n{doc}")
        return (
            "以下片段来自本地向量知识库（含已保存记忆与已入库文件），供回答参考：\n\n"
            + "\n\n---\n\n".join(parts)
        )
